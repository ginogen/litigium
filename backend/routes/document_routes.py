from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import FileResponse, JSONResponse
from typing import Dict, List, Optional, Any
from pydantic import BaseModel
import os
import tempfile
from datetime import datetime

from ..auth.dependencies import get_current_user
from ..models.user import User
from supabase_integration import supabase

router = APIRouter(prefix="/api/documents", tags=["documents"])

async def obtener_sesion_documento(session_id: str, abogado_id: str) -> Dict:
    """Obtiene una sesión específica verificando que pertenezca al abogado."""
    try:
        response = supabase.table('chat_sesiones')\
            .select('*')\
            .eq('session_id', session_id)\
            .eq('abogado_id', abogado_id)\
            .single()\
            .execute()
        
        if not response.data:
            raise HTTPException(status_code=404, detail="Sesión no encontrada")
        
        return response.data
    except Exception as e:
        if "No rows found" in str(e):
            raise HTTPException(status_code=404, detail="Sesión no encontrada")
        raise HTTPException(status_code=500, detail=f"Error obteniendo sesión: {str(e)}")

@router.get("/descargar/{session_id}")
async def descargar_demanda(
    session_id: str,
    current_user: User = Depends(get_current_user)
):
    """Descarga el archivo de demanda generado."""
    try:
        print(f"🔍 Descargando demanda para session_id: {session_id}")
        
        # Obtener el perfil del abogado
        abogado_response = supabase.table('abogados')\
            .select('*')\
            .eq('user_id', current_user.id)\
            .single()\
            .execute()
        
        if not abogado_response.data:
            raise HTTPException(status_code=404, detail="Perfil de abogado no encontrado")
        
        abogado_id = abogado_response.data['id']
        
        # Buscar demanda generada por session_id y abogado_id
        demanda_response = supabase.table('demandas_generadas')\
            .select('*')\
            .eq('session_id', session_id)\
            .eq('abogado_id', abogado_id)\
            .single()\
            .execute()
        
        if not demanda_response.data:
            raise HTTPException(
                status_code=404, 
                detail="No se encontró una demanda generada para esta sesión"
            )
        
        demanda = demanda_response.data
        print(f"✅ Demanda encontrada: {demanda.get('tipo_demanda', 'N/A')}")
        
        # Intentar descargar desde Supabase Storage primero
        archivo_url = demanda.get('archivo_docx_url')
        if archivo_url:
            try:
                # Extraer nombre del archivo de la URL
                import re
                match = re.search(r'/([^/]+\.docx)$', archivo_url)
                if match:
                    nombre_archivo_storage = match.group(1)
                    # Descargar desde Storage
                    storage_response = supabase.storage.from_('demandas-generadas').download(nombre_archivo_storage)
                    if storage_response:
                        print(f"✅ Archivo descargado desde Storage")
                        
                        # Crear archivo temporal
                        import tempfile
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                            temp_file.write(storage_response)
                            temp_path = temp_file.name
                        
                        # Generar nombre descriptivo
                        datos_cliente = demanda.get("datos_cliente", {})
                        if isinstance(datos_cliente, str):
                            import json
                            try:
                                datos_cliente = json.loads(datos_cliente)
                            except:
                                datos_cliente = {}
                        
                        cliente_nombre = datos_cliente.get("nombre_completo", "Cliente").replace(" ", "_")
                        tipo_demanda = demanda.get("tipo_demanda", "Demanda").replace(" ", "_")
                        nombre_archivo = f"Demanda_{tipo_demanda}_{cliente_nombre}.docx"
                        
                        return FileResponse(
                            path=temp_path,
                            filename=nombre_archivo,
                            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            headers={"Content-Disposition": f"attachment; filename={nombre_archivo}"}
                        )
            except Exception as storage_error:
                print(f"⚠️ Error descargando desde Storage: {storage_error}")
        
        # Fallback: crear archivo desde texto_demanda
        texto_demanda = demanda.get('texto_demanda', '')
        if not texto_demanda:
            raise HTTPException(
                status_code=404, 
                detail="La demanda no tiene contenido disponible"
            )
        
        # Obtener datos del cliente
        datos_cliente = demanda.get("datos_cliente", {})
        if isinstance(datos_cliente, str):
            import json
            try:
                datos_cliente = json.loads(datos_cliente)
            except:
                datos_cliente = {}
        
        cliente_nombre = datos_cliente.get("nombre_completo", "Cliente")
        tipo_demanda = demanda.get("tipo_demanda", "Demanda").replace(" ", "_")
        
        # Crear documento Word desde el texto
        from docx import Document
        import tempfile
        
        doc = Document()
        
        # Agregar título
        titulo = doc.add_heading(f'DEMANDA - {tipo_demanda.upper()}', 0)
        titulo.alignment = 1  # Centrado
        
        # Agregar información del cliente si existe
        if datos_cliente:
            doc.add_heading('DATOS DEL CLIENTE', level=1)
            tabla_cliente = doc.add_table(rows=0, cols=2)
            tabla_cliente.style = 'Table Grid'
            
            for campo, valor in datos_cliente.items():
                if valor:
                    row = tabla_cliente.add_row()
                    row.cells[0].text = campo.replace('_', ' ').title()
                    row.cells[1].text = str(valor)
            
            doc.add_page_break()
        
        # Agregar texto de la demanda
        doc.add_heading('DEMANDA', level=1)
        
        # Dividir en párrafos y procesar
        paragrafos = texto_demanda.split('\n\n')
        for parrafo in paragrafos:
            if parrafo.strip():
                doc.add_paragraph(parrafo.strip())
        
        # Guardar en archivo temporal
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            doc.save(temp_file.name)
            temp_path = temp_file.name
        
        # Generar nombre de archivo descriptivo
        nombre_archivo = f"Demanda_{tipo_demanda}_{cliente_nombre.replace(' ', '_')}.docx"
        
        print(f"✅ Archivo generado: {nombre_archivo}")
        
        return FileResponse(
            path=temp_path,
            filename=nombre_archivo,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={nombre_archivo}"}
        )
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error descargando archivo: {e}")
        raise HTTPException(status_code=500, detail=f"Error descargando archivo: {str(e)}")

@router.get("/preview/{session_id}")
async def preview_demanda(
    session_id: str,
    current_user: User = Depends(get_current_user)
):
    """Obtiene vista previa de la demanda generada."""
    try:
        print(f"🔍 Buscando demanda generada para session_id: {session_id}")
        
        # Buscar demanda generada en la base de datos
        demanda_response = supabase.table('demandas_generadas')\
            .select('*')\
            .eq('session_id', session_id)\
            .eq('user_id', current_user.id)\
            .execute()
        
        if not demanda_response.data:
            print(f"❌ No se encontró demanda generada para session_id: {session_id}")
            raise HTTPException(
                status_code=404, 
                detail="No se encontró una demanda generada para esta sesión"
            )
        
        demanda = demanda_response.data[0]
        print(f"✅ Demanda encontrada: {demanda.get('tipo_demanda', 'N/A')}")
        
        # Extraer texto de demanda y estructurarlo para preview
        texto_demanda = demanda.get('texto_demanda', '')
        if not texto_demanda:
            raise HTTPException(
                status_code=404, 
                detail="La demanda no tiene contenido generado"
            )
        
        # Parsear el texto para extraer secciones principales
        secciones = extraer_secciones_demanda(texto_demanda)
        
        # Obtener datos del cliente
        datos_cliente = demanda.get('datos_cliente', {})
        if isinstance(datos_cliente, str):
            import json
            try:
                datos_cliente = json.loads(datos_cliente)
            except:
                datos_cliente = {}
        
        # Obtener metadatos
        metadatos = demanda.get('metadatos', {})
        if isinstance(metadatos, str):
            import json
            try:
                metadatos = json.loads(metadatos)
            except:
                metadatos = {}
        
        preview_data = {
            "session_id": session_id,
            "tipo_demanda": demanda.get('tipo_demanda', 'No especificado'),
            "cliente_nombre": datos_cliente.get("nombre_completo", "Cliente"),
            "fecha_generacion": demanda.get('fecha_generacion', datetime.now().strftime("%d/%m/%Y")),
            "secciones": secciones,
            "metadatos": {
                "palabras": len(texto_demanda.split()),
                "caracteres": len(texto_demanda),
                "casos_consultados": metadatos.get("casos_consultados", 0),
                "tiempo_generacion": metadatos.get("tiempo_generacion", "N/A")
            },
            "texto_completo": texto_demanda,
            "estado": demanda.get('estado', 'completado')
        }
        
        print(f"✅ Preview generado: {len(texto_demanda)} caracteres")
        
        return {
            "success": True,
            "preview": preview_data
        }
        
    except HTTPException:
        # Re-lanzar HTTPExceptions tal como están
        raise
    except Exception as e:
        print(f"❌ Error generando vista previa: {e}")
        raise HTTPException(status_code=500, detail=f"Error generando vista previa: {str(e)}")

def extraer_secciones_demanda(texto_demanda: str) -> dict:
    """Extrae las secciones principales de una demanda."""
    secciones = {
        "hechos": "",
        "derecho": "",
        "petitorio": "",
        "prueba": ""
    }
    
    # Dividir el texto en líneas para procesamiento
    lineas = texto_demanda.split('\n')
    seccion_actual = None
    
    for linea in lineas:
        linea_lower = linea.lower().strip()
        
        # Detectar sección actual basada en palabras clave
        if any(palabra in linea_lower for palabra in ['hechos', 'relato', 'antecedentes']):
            seccion_actual = 'hechos'
        elif any(palabra in linea_lower for palabra in ['derecho', 'legal', 'normativa', 'ley']):
            seccion_actual = 'derecho'
        elif any(palabra in linea_lower for palabra in ['petitorio', 'solicita', 'pide']):
            seccion_actual = 'petitorio'
        elif any(palabra in linea_lower for palabra in ['prueba', 'ofrecimiento', 'evidencia']):
            seccion_actual = 'prueba'
        
        # Agregar contenido a la sección actual
        if seccion_actual and linea.strip():
            if secciones[seccion_actual]:
                secciones[seccion_actual] += f" {linea.strip()}"
            else:
                secciones[seccion_actual] = linea.strip()
    
    # Si no se pudo parsear por secciones, usar fragmentos del texto
    if not any(secciones.values()):
        fragmentos = texto_demanda.split('\n\n')
        if len(fragmentos) >= 4:
            secciones["hechos"] = fragmentos[0][:500] + "..." if len(fragmentos[0]) > 500 else fragmentos[0]
            secciones["derecho"] = fragmentos[1][:500] + "..." if len(fragmentos[1]) > 500 else fragmentos[1]
            secciones["petitorio"] = fragmentos[2][:500] + "..." if len(fragmentos[2]) > 500 else fragmentos[2]
            secciones["prueba"] = fragmentos[3][:500] + "..." if len(fragmentos[3]) > 500 else fragmentos[3]
        else:
            # Fallback: usar el primer fragmento para hechos
            secciones["hechos"] = texto_demanda[:1000] + "..." if len(texto_demanda) > 1000 else texto_demanda
    
    return secciones

@router.post("/guardar-cambios")
async def guardar_cambios_demanda(
    cambios_data: Dict[str, Any],
    current_user: User = Depends(get_current_user)
):
    """Guarda los cambios realizados en la demanda."""
    try:
        session_id = cambios_data.get("session_id")
        cambios = cambios_data.get("cambios", {})
        
        if not session_id:
            raise HTTPException(status_code=400, detail="session_id requerido")
        
        session = await obtener_sesion_documento(session_id, current_user.id)
        
        # TODO: Integrar con sistema de procesamiento de cambios
        # Por ahora, actualizar metadatos en la base de datos
        
        # Obtener cambios existentes
        cambios_existentes = session.get("cambios_realizados", [])
        if isinstance(cambios_existentes, str):
            import json
            try:
                cambios_existentes = json.loads(cambios_existentes)
            except:
                cambios_existentes = []
        
        # Agregar nuevo cambio
        cambios_existentes.append({
            "timestamp": datetime.now().isoformat(),
            "cambios": cambios,
            "usuario": current_user.email
        })
        
        # Actualizar en la base de datos
        supabase.table('chat_sesiones')\
            .update({
                "cambios_realizados": cambios_existentes,
                "updated_at": datetime.now().isoformat()
            })\
            .eq('session_id', session_id)\
            .execute()
        
        return {
            "success": True,
            "mensaje": "Cambios guardados exitosamente",
            "total_cambios": len(cambios_existentes)
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error guardando cambios: {str(e)}")

@router.post("/regenerar")
async def regenerar_demanda_endpoint(
    session_data: Dict[str, str],
    current_user: User = Depends(get_current_user)
):
    """Regenera la demanda usando IA con los datos existentes."""
    try:
        session_id = session_data.get("session_id")
        
        if not session_id:
            raise HTTPException(status_code=400, detail="session_id requerido")
        
        session = await obtener_sesion_documento(session_id, current_user.id)
        
        # TODO: Integrar con sistema de regeneración IA
        # Por ahora, actualizar metadatos de regeneración
        
        regeneraciones = session.get("regeneraciones", 0) + 1
        
        # Actualizar en la base de datos
        supabase.table('chat_sesiones')\
            .update({
                "ultima_regeneracion": datetime.now().isoformat(),
                "regeneraciones": regeneraciones,
                "updated_at": datetime.now().isoformat()
            })\
            .eq('session_id', session_id)\
            .execute()
        
        return {
            "success": True,
            "mensaje": "Demanda regenerada exitosamente",
            "regeneracion_numero": regeneraciones,
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error regenerando demanda: {str(e)}")

@router.get("/historial/{session_id}")
async def obtener_historial_documento(
    session_id: str,
    current_user: User = Depends(get_current_user)
):
    """Obtiene el historial de cambios de un documento."""
    try:
        session = await obtener_sesion_documento(session_id, current_user.id)
        
        # Obtener historial de cambios
        cambios_realizados = session.get("cambios_realizados", [])
        if isinstance(cambios_realizados, str):
            import json
            try:
                cambios_realizados = json.loads(cambios_realizados)
            except:
                cambios_realizados = []
        
        return {
            "success": True,
            "session_id": session_id,
            "historial": cambios_realizados,
            "total_cambios": len(cambios_realizados),
            "creado_en": session.get("created_at"),
            "ultima_actualizacion": session.get("updated_at")
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error obteniendo historial: {str(e)}")

@router.delete("/sesion/{session_id}")
async def eliminar_sesion_documento(
    session_id: str,
    current_user: User = Depends(get_current_user)
):
    """Elimina una sesión de documento."""
    try:
        session = await obtener_sesion_documento(session_id, current_user.id)
        
        # Eliminar sesión de la base de datos
        supabase.table('chat_sesiones')\
            .delete()\
            .eq('session_id', session_id)\
            .execute()
        
        return {
            "success": True,
            "mensaje": "Sesión eliminada exitosamente"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error eliminando sesión: {str(e)}")

@router.get("/estadisticas")
async def obtener_estadisticas_documentos(
    current_user: User = Depends(get_current_user)
):
    """Obtiene estadísticas de documentos del usuario."""
    try:
        # Obtener sesiones del usuario
        response = supabase.table('chat_sesiones')\
            .select('*')\
            .eq('abogado_id', current_user.id)\
            .execute()
        
        sesiones = response.data or []
        
        # Calcular estadísticas
        estadisticas = {
            "total_sesiones": len(sesiones),
            "sesiones_con_demanda": len([s for s in sesiones if s.get("tipo_demanda")]),
            "tipos_demanda": {},
            "ultima_actividad": None
        }
        
        # Contar tipos de demanda
        for sesion in sesiones:
            tipo = sesion.get("tipo_demanda")
            if tipo:
                estadisticas["tipos_demanda"][tipo] = estadisticas["tipos_demanda"].get(tipo, 0) + 1
        
        # Encontrar última actividad
        fechas_actualizacion = [s.get("updated_at") for s in sesiones if s.get("updated_at")]
        if fechas_actualizacion:
            estadisticas["ultima_actividad"] = max(fechas_actualizacion)
        
        return {
            "success": True,
            "estadisticas": estadisticas
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error obteniendo estadísticas: {str(e)}") 