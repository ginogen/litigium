import os
import uuid
import hashlib
from typing import Dict, List, Optional, Tuple, Any
from pathlib import Path
import tempfile
import mimetypes

# Document processing
from docx import Document
import PyPDF2

# Vector embeddings and storage
from qdrant_client import QdrantClient
from qdrant_client.models import PointStruct, Distance, VectorParams, Filter, FieldCondition, MatchValue
from langchain_openai import OpenAIEmbeddings

# Environment and database
from dotenv import load_dotenv
from supabase import create_client, Client

# Importar el procesador de formato rico
try:
    from .document_processor_rich_format import rich_format_processor
except ImportError:
    print("⚠️ Procesador de formato rico no disponible, usando solo texto plano")
    rich_format_processor = None

load_dotenv()

class DocumentProcessor:
    """Procesador de documentos para usuarios específicos con Qdrant personalizado."""
    
    def __init__(self):
        # Configuración de Qdrant
        self.qdrant_url = os.getenv("QDRANT_URL")
        self.qdrant_api_key = os.getenv("QDRANT_API_KEY")
        self.collection_prefix = os.getenv("QDRANT_COLLECTION_PREFIX", "legal_")
        
        # Configuración de OpenAI
        self.openai_api_key = os.getenv("OPENAI_API_KEY")
        
        # Configuración de Supabase
        self.supabase_url = os.getenv("SUPABASE_URL")
        self.supabase_key = os.getenv("SUPABASE_ANON_KEY")
        
        # Inicializar clientes
        self.qdrant_client = QdrantClient(
            url=self.qdrant_url, 
            api_key=self.qdrant_api_key
        )
        self.embedder = OpenAIEmbeddings(openai_api_key=self.openai_api_key)
        self.supabase: Client = create_client(self.supabase_url, self.supabase_key)
    
    def collection_exists(self, collection_name: str) -> bool:
        """Verifica si una colección existe en Qdrant."""
        try:
            collections = self.qdrant_client.get_collections()
            collection_names = [collection.name for collection in collections.collections]
            return collection_name in collection_names
        except Exception as e:
            print(f"❌ Error verificando colección {collection_name}: {e}")
            return False

    def get_user_collection_name(self, user_id: str) -> str:
        """Genera el nombre de colección único para un usuario."""
        return f"{self.collection_prefix}user_{user_id}"
    
    def create_user_collection_if_not_exists(self, user_id: str) -> str:
        """Crea la colección personal del usuario en Qdrant si no existe."""
        collection_name = self.get_user_collection_name(user_id)
        
        try:
            if not self.collection_exists(collection_name):
                self.qdrant_client.create_collection(
                    collection_name=collection_name,
                    vectors_config=VectorParams(size=1536, distance=Distance.COSINE)
                )
                print(f"✅ Colección personal creada: {collection_name}")
            return collection_name
        except Exception as e:
            raise Exception(f"Error creando colección para usuario {user_id}: {e}")
    
    def extract_text_from_file(self, file_path: str, mime_type: str = None) -> Tuple[str, Dict]:
        """Extrae texto de diferentes tipos de archivo."""
        try:
            if mime_type is None:
                mime_type, _ = mimetypes.guess_type(file_path)
            
            text = ""
            metadata = {"extraction_method": ""}
            
            if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                # DOCX
                text, metadata = self._extract_from_docx(file_path)
            elif mime_type == "application/msword":
                # DOC - Intentar métodos alternativos
                text, metadata = self._extract_from_doc(file_path)
            elif mime_type == "application/pdf":
                # PDF
                text, metadata = self._extract_from_pdf(file_path)
            elif mime_type == "text/plain":
                # TXT
                text, metadata = self._extract_from_txt(file_path)
            else:
                raise Exception(f"Tipo de archivo no soportado: {mime_type}")
            
            return text.strip(), metadata
            
        except Exception as e:
            raise Exception(f"Error extrayendo texto del archivo: {e}")
    
    def _extract_from_docx(self, file_path: str) -> Tuple[str, Dict]:
        """Extrae texto de archivo DOCX."""
        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
            metadata = {
                "extraction_method": "python-docx",
                "paragraphs_count": len(doc.paragraphs)
            }
            return text, metadata
        except Exception as e:
            raise Exception(f"Error procesando DOCX: {e}")
    
    def _extract_from_doc(self, file_path: str) -> Tuple[str, Dict]:
        """Extrae texto de archivo DOC antiguo usando múltiples métodos."""
        methods = [
            ("python-docx", self._try_docx_on_doc),
            ("manual_conversion", self._try_manual_doc_conversion),
            ("textract", self._try_textract_extraction)
        ]
        
        for method_name, method_func in methods:
            try:
                text = method_func(file_path)
                if text and text.strip():
                    metadata = {
                        "extraction_method": method_name,
                        "fallback_used": True,
                        "file_size": os.path.getsize(file_path)
                    }
                    print(f"  ✅ DOC procesado con {method_name}")
                    return text, metadata
            except Exception as e:
                print(f"  ⚠️ Método {method_name} falló: {str(e)[:100]}")
                continue
        
        # Si ningún método funcionó, extraer texto básico
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
            
            # Intentar decodificar y extraer texto visible
            text = content.decode('latin-1', errors='ignore')
            # Limpiar caracteres no imprimibles
            clean_text = ''.join(char for char in text if char.isprintable() or char.isspace())
            
            if len(clean_text) > 100:  # Solo si hay contenido mínimo
                metadata = {
                    "extraction_method": "raw_decode_fallback",
                    "warning": "Extracción básica, puede contener caracteres extraños"
                }
                print(f"  ⚠️ DOC procesado con método básico")
                return clean_text, metadata  # Sin límite de caracteres para preservar contenido completo
        except:
            pass
        
        raise Exception("No se pudo extraer texto del archivo .doc. Convierte a .docx e inténtalo nuevamente.")
    
    def _try_docx_on_doc(self, file_path: str) -> str:
        """Intenta leer un .doc como si fuera .docx (a veces funciona)."""
        doc = Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
    
    def _try_textract_extraction(self, file_path: str) -> str:
        """Intenta usar textract para extraer texto."""
        try:
            import textract
            text = textract.process(file_path).decode('utf-8', errors='ignore')
            return text
        except ImportError:
            raise Exception("textract no está instalado")
    
    def _try_manual_doc_conversion(self, file_path: str) -> str:
        """Intenta convertir manualmente usando python-docx2txt."""
        try:
            import docx2txt
            text = docx2txt.process(file_path)
            return text if text else ""
        except ImportError:
            raise Exception("docx2txt no está instalado")
    
    def _extract_from_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """Extrae texto de archivo PDF."""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                pages_count = len(pdf_reader.pages)
                
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            metadata = {
                "extraction_method": "PyPDF2",
                "pages_count": pages_count
            }
            return text, metadata
        except Exception as e:
            raise Exception(f"Error procesando PDF: {e}")
    
    def _extract_from_txt(self, file_path: str) -> Tuple[str, Dict]:
        """Extrae texto de archivo TXT."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            metadata = {
                "extraction_method": "direct-read",
                "encoding": "utf-8"
            }
            return text, metadata
        except Exception as e:
            raise Exception(f"Error procesando TXT: {e}")
    
    def extract_legal_sections(self, text: str) -> Dict[str, str]:
        """Extrae secciones legales del texto (hechos, derecho, petitorio, etc.)."""
        sections = {
            "hechos": "",
            "derecho": "",
            "petitorio": "",
            "prueba": "",
            "conclusion": ""
        }
        
        current_section = None
        lines = text.split("\n")
        
        for line in lines:
            line_lower = line.strip().lower()
            
            # Detectar secciones
            if any(keyword in line_lower for keyword in ["hechos", "i.-", "i)", "1.-", "1)", "primera"]):
                current_section = "hechos"
            elif any(keyword in line_lower for keyword in ["derecho", "fundamento", "ii.-", "ii)", "2.-", "2)", "segunda"]):
                current_section = "derecho"
            elif any(keyword in line_lower for keyword in ["petitorio", "solicito", "pido", "iii.-", "iii)", "3.-", "3)", "tercera"]):
                current_section = "petitorio"
            elif any(keyword in line_lower for keyword in ["prueba", "ofrezco", "iv.-", "iv)", "4.-", "4)", "cuarta"]):
                current_section = "prueba"
            elif any(keyword in line_lower for keyword in ["conclusion", "conclusión", "v.-", "v)", "5.-", "5)", "quinta"]):
                current_section = "conclusion"
            elif current_section and line.strip():
                sections[current_section] += line.strip() + " "
        
        return sections
    
    def generate_embedding(self, text: str) -> List[float]:
        """Genera embedding para el texto dado."""
        try:
            return self.embedder.embed_query(text)
        except Exception as e:
            raise Exception(f"Error generando embedding: {e}")
    
    def process_and_upload_document(
        self, 
        user_id: str, 
        file_path: str, 
        filename: str,
        categoria_id: str,
        tipo_demanda: str,
        mime_type: str = None
    ) -> Dict[str, Any]:
        """Procesa un documento y lo sube a la colección personal del usuario en Qdrant."""
        
        try:
            # 1. Crear colección del usuario si no existe
            collection_name = self.create_user_collection_if_not_exists(user_id)
            
            # 2. Extraer texto del documento (método tradicional)
            texto_extraido, extraction_metadata = self.extract_text_from_file(file_path, mime_type)
            
            if not texto_extraido.strip():
                raise Exception("No se pudo extraer texto del documento")
            
            # 3. NUEVO: Extraer formato rico si está disponible
            contenido_rico = None
            metadatos_formato = None
            tiene_formato_rico = False
            
            if rich_format_processor and mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                try:
                    print("🎨 Extrayendo formato rico...")
                    contenido_rico, metadatos_formato = rich_format_processor.extract_rich_content_from_file(file_path, mime_type)
                    tiene_formato_rico = True
                    print(f"✅ Formato rico extraído: {metadatos_formato.get('total_blocks', 0)} bloques")
                except Exception as e:
                    print(f"⚠️ Error extrayendo formato rico: {e}")
                    contenido_rico = None
                    metadatos_formato = {"error": str(e), "fallback_to_plain": True}
            
            # 4. Extraer secciones legales
            secciones = self.extract_legal_sections(texto_extraido)
            
            # 5. Crear texto completo para embedding
            texto_completo = " ".join([
                secciones.get("hechos", ""),
                secciones.get("derecho", ""),
                secciones.get("petitorio", ""),
                secciones.get("prueba", ""),
                secciones.get("conclusion", "")
            ]).strip()
            
            # 6. Generar embedding
            vector = self.generate_embedding(texto_completo)
            
            # 7. Crear ID único para el punto
            doc_id = str(uuid.uuid4())
            
            # 8. Preparar payload extendido con formato rico
            payload = {
                "document_id": doc_id,
                "user_id": user_id,
                "filename": filename,
                "categoria_id": categoria_id,
                "tipo_demanda": tipo_demanda,
                "texto_completo": texto_completo,
                "secciones": secciones,
                "extraction_metadata": extraction_metadata,
                "processed_at": str(uuid.uuid1().time),
                "mime_type": mime_type,
                # NUEVO: Metadatos de formato rico
                "tiene_formato_rico": tiene_formato_rico,
                "metadatos_formato": metadatos_formato if metadatos_formato else {},
                "version_extraccion": "v2.0" if tiene_formato_rico else "v1.0"
            }
            
            # 8. Crear punto y subir a Qdrant
            punto = PointStruct(
                id=doc_id,
                vector=vector,
                payload=payload
            )
            
            self.qdrant_client.upsert(
                collection_name=collection_name,
                points=[punto]
            )
            
            return {
                "success": True,
                "document_id": doc_id,
                "collection_name": collection_name,
                "texto_extraido": texto_extraido,
                "secciones": secciones,
                "extraction_metadata": extraction_metadata,
                # NUEVO: Información de formato rico
                "tiene_formato_rico": tiene_formato_rico,
                "metadatos_formato": metadatos_formato,
                "contenido_rico": contenido_rico,
                "version_extraccion": "v2.0" if tiene_formato_rico else "v1.0"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    def search_similar_documents(
        self, 
        user_id: str, 
        query_text: str, 
        categoria_id: str = None,
        limit: int = 5
    ) -> List[Dict[str, Any]]:
        """Busca documentos similares en la colección personal del usuario."""
        
        try:
            collection_name = self.get_user_collection_name(user_id)
            
            # Verificar que la colección existe
            if not self.collection_exists(collection_name):
                return []
            
            # Generar embedding para la consulta
            query_vector = self.generate_embedding(query_text)
            
            # Preparar filtros si se especifica categoría
            query_filter = None
            if categoria_id:
                query_filter = Filter(
                    must=[
                        FieldCondition(
                            key="categoria_id",
                            match=MatchValue(value=categoria_id)
                        )
                    ]
                )
            
            # Realizar búsqueda
            results = self.qdrant_client.search(
                collection_name=collection_name,
                query_vector=query_vector,
                query_filter=query_filter,
                limit=limit,
                with_payload=True
            )
            
            # Formatear resultados
            documentos_similares = []
            for result in results:
                documentos_similares.append({
                    "id": result.id,
                    "score": result.score,
                    "filename": result.payload.get("filename"),
                    "tipo_demanda": result.payload.get("tipo_demanda"),
                    "categoria_id": result.payload.get("categoria_id"),
                    "texto_relevante": result.payload.get("texto_completo", "")[:500] + "...",
                    "secciones": result.payload.get("secciones", {})
                })
            
            return documentos_similares
            
        except Exception as e:
            raise Exception(f"Error buscando documentos similares: {e}")
    
    def get_user_collection_stats(self, user_id: str) -> Dict[str, Any]:
        """Obtiene estadísticas de la colección personal del usuario."""
        try:
            collection_name = self.get_user_collection_name(user_id)
            
            if not self.collection_exists(collection_name):
                return {
                    "total_documents": 0,
                    "collection_exists": False,
                    "collection_name": collection_name
                }
            
            # Usar scroll para contar documentos en lugar de get_collection 
            # para evitar errores de validación de Pydantic con versiones nuevas de Qdrant
            try:
                # Contar usando scroll con limit pequeño para obtener total
                scroll_result = self.qdrant_client.scroll(
                    collection_name=collection_name,
                    limit=1,
                    with_payload=False,
                    with_vectors=False
                )
                
                # Si hay al menos un documento, usar el método alternativo de conteo
                if scroll_result[0]:  # Si hay puntos
                    # Hacer scroll completo para contar (alternativa más segura)
                    all_points = self.qdrant_client.scroll(
                        collection_name=collection_name,
                        limit=10000,  # Límite alto para obtener todos
                        with_payload=False,
                        with_vectors=False
                    )
                    total_count = len(all_points[0])
                else:
                    total_count = 0
                
                return {
                    "total_documents": total_count,
                    "collection_exists": True,
                    "collection_name": collection_name
                }
                
            except Exception as count_error:
                print(f"⚠️ Error contando documentos con scroll, usando método fallback: {count_error}")
                # Fallback: intentar get_collection pero manejar errores de Pydantic
                try:
                    collection_info = self.qdrant_client.get_collection(collection_name)
                    points_count = getattr(collection_info, 'points_count', 0)
                    return {
                        "total_documents": points_count,
                        "collection_exists": True,
                        "collection_name": collection_name
                    }
                except Exception as pydantic_error:
                    print(f"⚠️ Error de validación Pydantic (ignorando): {pydantic_error}")
                    # Retornar estimación básica
                    return {
                        "total_documents": 0,  # No podemos contar exactamente
                        "collection_exists": True,
                        "collection_name": collection_name,
                        "warning": "No se pudo obtener conteo exacto debido a incompatibilidad de versiones"
                    }
            
        except Exception as e:
            raise Exception(f"Error obteniendo estadísticas: {e}")
    
    def delete_document_from_collection(self, user_id: str, document_id: str) -> bool:
        """Elimina un documento específico de la colección del usuario."""
        try:
            collection_name = self.get_user_collection_name(user_id)
            
            if not self.collection_exists(collection_name):
                return False
            
            self.qdrant_client.delete(
                collection_name=collection_name,
                points_selector=[document_id]
            )
            
            return True
            
        except Exception as e:
            raise Exception(f"Error eliminando documento: {e}")
    
    # ================== FUNCIONES DE ANOTACIONES ==================
    
    def update_document_with_annotations(
        self, 
        user_id: str, 
        document_id: str, 
        annotations: List[Dict[str, Any]],
        original_document_data: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """Actualiza un documento en Qdrant incluyendo el contexto de las anotaciones."""
        
        try:
            collection_name = self.get_user_collection_name(user_id)
            
            if not self.collection_exists(collection_name):
                raise Exception(f"Colección {collection_name} no existe")
            
            # Obtener documento original de Qdrant si no se proporciona
            if not original_document_data:
                original_document_data = self.get_document_from_collection(user_id, document_id)
                if not original_document_data:
                    raise Exception(f"Documento {document_id} no encontrado en Qdrant")
            
            # Construir contexto enriquecido con anotaciones
            enhanced_context = self._build_enhanced_context_with_annotations(
                original_document_data, 
                annotations
            )
            
            # Generar nuevo embedding con contexto enriquecido
            enhanced_vector = self.generate_embedding(enhanced_context["enhanced_text"])
            
            # Actualizar payload con información de anotaciones
            enhanced_payload = {
                **original_document_data.get("payload", {}),
                "annotations_context": enhanced_context["annotations_summary"],
                "annotations_count": len(annotations),
                "annotations_by_type": enhanced_context["annotations_by_type"],
                "expert_insights": enhanced_context["expert_insights"],
                "enhanced_strategies": enhanced_context["strategies"],
                "identified_precedents": enhanced_context["precedents"],
                "detected_problems": enhanced_context["problems"],
                "last_annotation_update": str(uuid.uuid1().time),
                "has_expert_annotations": True
            }
            
            # Actualizar punto en Qdrant
            updated_point = PointStruct(
                id=document_id,
                vector=enhanced_vector,
                payload=enhanced_payload
            )
            
            self.qdrant_client.upsert(
                collection_name=collection_name,
                points=[updated_point]
            )
            
            return {
                "success": True,
                "document_id": document_id,
                "annotations_processed": len(annotations),
                "enhanced_context": enhanced_context,
                "collection_name": collection_name
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    def _build_enhanced_context_with_annotations(
        self, 
        original_document: Dict[str, Any], 
        annotations: List[Dict[str, Any]]
    ) -> Dict[str, Any]:
        """Construye contexto enriquecido combinando documento original con anotaciones."""
        
        # Texto original
        original_text = original_document.get("payload", {}).get("texto_completo", "")
        secciones = original_document.get("payload", {}).get("secciones", {})
        
        # Categorizar anotaciones por tipo
        annotations_by_type = {
            "comentario": [],
            "precedente": [],
            "estrategia": [],
            "problema": []
        }
        
        for annotation in annotations:
            tipo = annotation.get("tipo_anotacion", "comentario")
            if tipo in annotations_by_type:
                annotations_by_type[tipo].append(annotation)
        
        # Construir contexto enriquecido
        enhanced_parts = [original_text]
        
        # Añadir insights de anotaciones
        if annotations_by_type["precedente"]:
            precedent_text = " ".join([
                f"PRECEDENTE IDENTIFICADO: {ann['contenido']}" 
                for ann in annotations_by_type["precedente"]
            ])
            enhanced_parts.append(precedent_text)
        
        if annotations_by_type["estrategia"]:
            strategy_text = " ".join([
                f"ESTRATEGIA LEGAL: {ann['contenido']}" 
                for ann in annotations_by_type["estrategia"]
            ])
            enhanced_parts.append(strategy_text)
        
        if annotations_by_type["problema"]:
            problem_text = " ".join([
                f"PROBLEMA DETECTADO: {ann['contenido']}" 
                for ann in annotations_by_type["problema"]
            ])
            enhanced_parts.append(problem_text)
        
        if annotations_by_type["comentario"]:
            comment_text = " ".join([
                f"OBSERVACIÓN EXPERTA: {ann['contenido']}" 
                for ann in annotations_by_type["comentario"]
            ])
            enhanced_parts.append(comment_text)
        
        # Texto completo enriquecido
        enhanced_text = " ".join(enhanced_parts)
        
        # Resumen de anotaciones para búsquedas
        annotations_summary = self._create_annotations_summary(annotations)
        
        return {
            "enhanced_text": enhanced_text,
            "annotations_summary": annotations_summary,
            "annotations_by_type": {k: len(v) for k, v in annotations_by_type.items()},
            "expert_insights": [ann["contenido"] for ann in annotations if ann.get("tipo_anotacion") == "comentario"],
            "strategies": [ann["contenido"] for ann in annotations if ann.get("tipo_anotacion") == "estrategia"],
            "precedents": [ann["contenido"] for ann in annotations if ann.get("tipo_anotacion") == "precedente"],
            "problems": [ann["contenido"] for ann in annotations if ann.get("tipo_anotacion") == "problema"]
        }
    
    def _create_annotations_summary(self, annotations: List[Dict[str, Any]]) -> str:
        """Crea un resumen textual de las anotaciones para búsquedas."""
        
        if not annotations:
            return ""
        
        summary_parts = []
        
        # Agrupar por tipo y crear resumen
        types_summary = {}
        for annotation in annotations:
            tipo = annotation.get("tipo_anotacion", "comentario")
            if tipo not in types_summary:
                types_summary[tipo] = []
            
            # Extraer palabras clave del contenido
            content = annotation.get("contenido", "")
            etiquetas = annotation.get("etiquetas", [])
            
            summary_parts.extend(etiquetas)
            summary_parts.append(content[:100])  # Primeras 100 chars
        
        return " ".join(summary_parts)
    
    def get_document_from_collection(self, user_id: str, document_id: str) -> Dict[str, Any]:
        """Obtiene un documento específico de la colección del usuario."""
        
        try:
            collection_name = self.get_user_collection_name(user_id)
            
            if not self.collection_exists(collection_name):
                return None
            
            # Buscar el documento por ID
            result = self.qdrant_client.retrieve(
                collection_name=collection_name,
                ids=[document_id],
                with_payload=True,
                with_vectors=True
            )
            
            if result and len(result) > 0:
                point = result[0]
                return {
                    "id": point.id,
                    "vector": point.vector,
                    "payload": point.payload
                }
            
            return None
            
        except Exception as e:
            raise Exception(f"Error obteniendo documento: {e}")
    
    def search_with_annotations_context(
        self, 
        user_id: str, 
        query_text: str, 
        categoria_id: str = None,
        include_annotations: bool = True,
        limit: int = 5
    ) -> List[Dict[str, Any]]:
        """Busca documentos considerando el contexto de anotaciones."""
        
        try:
            collection_name = self.get_user_collection_name(user_id)
            
            if not self.collection_exists(collection_name):
                return []
            
            # Generar embedding para la consulta
            query_vector = self.generate_embedding(query_text)
            
            # Preparar filtros
            query_filter = None
            filter_conditions = []
            
            if categoria_id:
                filter_conditions.append(
                    FieldCondition(
                        key="categoria_id",
                        match=MatchValue(value=categoria_id)
                    )
                )
            
            # Priorizar documentos con anotaciones si se solicita
            if include_annotations:
                filter_conditions.append(
                    FieldCondition(
                        key="has_expert_annotations",
                        match=MatchValue(value=True)
                    )
                )
            
            if filter_conditions:
                query_filter = Filter(must=filter_conditions)
            
            # Realizar búsqueda
            results = self.qdrant_client.search(
                collection_name=collection_name,
                query_vector=query_vector,
                query_filter=query_filter,
                limit=limit,
                with_payload=True
            )
            
            # Si no encuentra suficientes con anotaciones, buscar sin filtro
            if include_annotations and len(results) < limit // 2:
                additional_results = self.qdrant_client.search(
                    collection_name=collection_name,
                    query_vector=query_vector,
                    query_filter=Filter(
                        must=[FieldCondition(key="categoria_id", match=MatchValue(value=categoria_id))]
                    ) if categoria_id else None,
                    limit=limit - len(results),
                    with_payload=True
                )
                results.extend(additional_results)
            
            # Formatear resultados con contexto de anotaciones
            documentos_enriquecidos = []
            for result in results:
                payload = result.payload
                
                documento = {
                    "id": result.id,
                    "score": result.score,
                    "filename": payload.get("filename"),
                    "tipo_demanda": payload.get("tipo_demanda"),
                    "categoria_id": payload.get("categoria_id"),
                    "texto_relevante": payload.get("texto_completo", "")[:500] + "...",
                    "secciones": payload.get("secciones", {}),
                    "has_annotations": payload.get("has_expert_annotations", False),
                    "annotations_count": payload.get("annotations_count", 0)
                }
                
                # Añadir contexto de anotaciones si existe
                if payload.get("has_expert_annotations"):
                    documento.update({
                        "expert_insights": payload.get("expert_insights", []),
                        "strategies": payload.get("enhanced_strategies", []),
                        "precedents": payload.get("identified_precedents", []),
                        "problems": payload.get("detected_problems", []),
                        "annotations_by_type": payload.get("annotations_by_type", {})
                    })
                
                documentos_enriquecidos.append(documento)
            
            return documentos_enriquecidos
            
        except Exception as e:
            raise Exception(f"Error buscando con contexto de anotaciones: {e}")
    
    def get_enhanced_legal_context(
        self, 
        user_id: str, 
        query_text: str, 
        tipo_demanda: str = None,
        limit: int = 3
    ) -> Dict[str, Any]:
        """Obtiene contexto legal enriquecido para generación de demandas."""
        
        try:
            # Buscar documentos con anotaciones
            documentos = self.search_with_annotations_context(
                user_id=user_id,
                query_text=query_text,
                include_annotations=True,
                limit=limit
            )
            
            if not documentos:
                return {
                    "contexto_textual": "No se encontraron casos similares en tu colección personal.",
                    "estrategias_expertas": [],
                    "precedentes_identificados": [],
                    "problemas_comunes": [],
                    "insights_adicionales": []
                }
            
            # Consolidar contexto enriquecido
            context_parts = []
            all_strategies = []
            all_precedents = []
            all_problems = []
            all_insights = []
            
            for i, doc in enumerate(documentos):
                # Texto base del documento
                context_parts.append(f"CASO {i+1} (Similaridad: {doc['score']:.2f}): {doc['texto_relevante']}")
                
                # Añadir contexto de anotaciones si existe
                if doc.get("has_annotations"):
                    if doc.get("strategies"):
                        strategies_text = " | ".join(doc["strategies"])
                        context_parts.append(f"ESTRATEGIAS IDENTIFICADAS: {strategies_text}")
                        all_strategies.extend(doc["strategies"])
                    
                    if doc.get("precedents"):
                        precedents_text = " | ".join(doc["precedents"])
                        context_parts.append(f"PRECEDENTES APLICABLES: {precedents_text}")
                        all_precedents.extend(doc["precedents"])
                    
                    if doc.get("problems"):
                        problems_text = " | ".join(doc["problems"])
                        context_parts.append(f"PROBLEMAS DETECTADOS: {problems_text}")
                        all_problems.extend(doc["problems"])
                    
                    if doc.get("expert_insights"):
                        insights_text = " | ".join(doc["expert_insights"])
                        context_parts.append(f"INSIGHTS EXPERTOS: {insights_text}")
                        all_insights.extend(doc["expert_insights"])
            
            return {
                "contexto_textual": " || ".join(context_parts),
                "estrategias_expertas": list(set(all_strategies)),  # Eliminar duplicados
                "precedentes_identificados": list(set(all_precedents)),
                "problemas_comunes": list(set(all_problems)),
                "insights_adicionales": list(set(all_insights)),
                "documentos_fuente": len(documentos),
                "documentos_con_anotaciones": len([d for d in documentos if d.get("has_annotations")])
            }
            
        except Exception as e:
            raise Exception(f"Error obteniendo contexto legal enriquecido: {e}")