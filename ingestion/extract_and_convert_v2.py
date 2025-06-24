from docx import Document
import json
import os
import re
from glob import glob
import zipfile
import tempfile
import shutil

def extract_sections(text):
    """Extrae las secciones principales de una demanda."""
    sections = {"hechos": "", "derecho": "", "petitorio": "", "prueba": ""}
    current = None
    
    for line in text.split("\n"):
        line_lower = line.strip().lower()
        if any(word in line_lower for word in ["hechos", "i.-", "i)", "1.-", "1)", "primero"]):
            current = "hechos"
        elif any(word in line_lower for word in ["derecho", "ii.-", "ii)", "2.-", "2)", "segundo"]):
            current = "derecho"
        elif any(word in line_lower for word in ["petitorio", "solicito", "iii.-", "iii)", "3.-", "3)", "tercero", "pido"]):
            current = "petitorio"
        elif any(word in line_lower for word in ["prueba", "ofrezco", "iv.-", "iv)", "4.-", "4)", "cuarto"]):
            current = "prueba"
        elif current and line.strip():
            sections[current] += line.strip() + " "
    
    return sections

def inferir_tipo_por_carpeta(filepath):
    """Infiere el tipo de demanda basado en la carpeta donde está el archivo."""
    carpeta = os.path.basename(os.path.dirname(filepath)).lower()
    
    tipos_mapeados = {
        "despido": "Despido Injustificado",
        "discriminatorio": "Despido Discriminatorio", 
        "accidente": "Accidente Laboral",
        "licencia": "Licencia Médica",
        "solidaridad": "Solidaridad Laboral",
        "blanco": "Empleados en Blanco",
        "negro": "Empleados en Negro"
    }
    
    for key, tipo in tipos_mapeados.items():
        if key in carpeta:
            return tipo
    
    return carpeta.replace("_", " ").title()

def es_archivo_word_valido(filepath):
    """Verifica si el archivo es un documento Word válido."""
    try:
        # Verificar si es un archivo ZIP válido (estructura de docx/doc moderna)
        with zipfile.ZipFile(filepath, 'r') as zip_file:
            # Verificar estructura básica de Word
            required_files = ['word/document.xml', '[Content_Types].xml']
            zip_contents = zip_file.namelist()
            
            for req_file in required_files:
                if req_file not in zip_contents:
                    return False
        return True
    except:
        return False

def intentar_reparar_archivo(filepath):
    """Intenta diferentes métodos para leer archivos problemáticos."""
    methods = [
        # Método 1: Leer directamente
        lambda f: Document(f),
        
        # Método 2: Copiar a temporal y leer
        lambda f: _leer_con_copia_temporal(f),
        
        # Método 3: Solo archivos .docx
        lambda f: Document(f) if f.endswith('.docx') else None
    ]
    
    for i, method in enumerate(methods):
        try:
            doc = method(filepath)
            if doc:
                print(f"  ✓ Método {i+1} exitoso")
                return doc
        except Exception as e:
            print(f"  ⚠️ Método {i+1} falló: {str(e)[:100]}")
            continue
    
    return None

def _leer_con_copia_temporal(filepath):
    """Copia el archivo a temporal e intenta leerlo."""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        shutil.copy2(filepath, tmp.name)
        try:
            return Document(tmp.name)
        finally:
            os.unlink(tmp.name)

def docx_to_json_mejorado(filepath, output_folder="data/casos_anteriores"):
    """Versión mejorada que maneja archivos problemáticos."""
    nombre_archivo = os.path.basename(filepath)
    print(f"🔄 Procesando: {nombre_archivo}")
    
    try:
        # Verificar si es un archivo Word válido
        if not es_archivo_word_valido(filepath):
            print(f"  ⚠️ Archivo no es Word válido, intentando métodos alternativos...")
        
        # Intentar múltiples métodos de lectura
        doc = intentar_reparar_archivo(filepath)
        
        if not doc:
            print(f"  ❌ No se pudo leer el archivo")
            return False
        
        # Extraer texto
        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        
        if not full_text.strip():
            print(f"  ⚠️ Archivo sin contenido de texto")
            return False

        # Procesar secciones
        sections = extract_sections(full_text)
        
        # Verificar que se extrajeron secciones
        has_content = any(sections[key].strip() for key in sections)
        if not has_content:
            print(f"  ⚠️ No se encontraron secciones legales reconocibles")
            # Guardar todo el texto en 'hechos' como fallback
            sections["hechos"] = full_text  # Preservar contenido completo sin limitaciones
        
        # Inferir tipo por carpeta
        tipo_inferido = inferir_tipo_por_carpeta(filepath)
        
        # Crear JSON
        json_data = {
            "id": os.path.splitext(nombre_archivo)[0],
            "tipo": tipo_inferido,
            "jurisdiccion": "Argentina",
            "empresa": "",
            "fecha": "",
            **sections,
            "archivo_origen": filepath,
            "procesado_exitosamente": True
        }

        # Guardar
        os.makedirs(output_folder, exist_ok=True)
        output_path = os.path.join(output_folder, json_data["id"] + ".json")
        
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)

        print(f"  ✅ Convertido → {tipo_inferido}")
        return True
        
    except Exception as e:
        print(f"  ❌ Error final: {str(e)[:100]}")
        return False

def procesar_archivos_mejorado():
    """Procesa archivos con mejor manejo de errores."""
    print("🚀 PROCESADOR MEJORADO DE DOCUMENTOS LEGALES")
    print("=" * 50)
    
    # Buscar archivos
    archivos_word = []
    for root, dirs, files in os.walk("ejemplos_demandas"):
        for file in files:
            if file.endswith(('.docx', '.doc')):
                archivos_word.append(os.path.join(root, file))
    
    if not archivos_word:
        print("⚠️ No se encontraron archivos Word")
        return
    
    print(f"📂 Encontrados: {len(archivos_word)} archivos")
    
    # Separar por extensión
    docx_files = [f for f in archivos_word if f.endswith('.docx')]
    doc_files = [f for f in archivos_word if f.endswith('.doc')]
    
    print(f"   • {len(docx_files)} archivos .docx")
    print(f"   • {len(doc_files)} archivos .doc")
    print()
    
    # Procesar archivos
    exitosos = 0
    errores = []
    
    for archivo in archivos_word:
        if docx_to_json_mejorado(archivo):
            exitosos += 1
        else:
            errores.append(os.path.basename(archivo))
    
    # Reporte final
    print("\n" + "="*50)
    print("📊 REPORTE FINAL")
    print(f"✅ Exitosos: {exitosos}/{len(archivos_word)}")
    print(f"❌ Errores: {len(errores)}")
    
    if errores:
        print(f"\n❌ Archivos problemáticos:")
        for error in errores[:10]:  # Mostrar solo los primeros 10
            print(f"   • {error}")
        if len(errores) > 10:
            print(f"   ... y {len(errores) - 10} más")
    
    print(f"\n📁 Archivos JSON guardados en: data/casos_anteriores/")
    
    return exitosos, len(errores)

if __name__ == "__main__":
    procesar_archivos_mejorado() 