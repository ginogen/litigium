import os
from typing import Dict, List, Optional
from dotenv import load_dotenv
from qdrant_client import QdrantClient
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain.prompts import ChatPromptTemplate
from docx import Document
from docx.shared import Inches
from datetime import datetime
import tempfile
import json
import time

load_dotenv()

# Variables de entorno
QDRANT_URL = os.getenv("QDRANT_URL")
QDRANT_API_KEY = os.getenv("QDRANT_API_KEY")
COLLECTION_PREFIX = os.getenv("QDRANT_COLLECTION_PREFIX", "legalbot_")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# Inicialización de servicios
qdrant = QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY)
llm = ChatOpenAI(
    model_name="gpt-4", 
    temperature=0.3, 
    openai_api_key=OPENAI_API_KEY,
    request_timeout=45,  # Timeout de 45 segundos
    max_retries=1  # Solo 1 reintento para no duplicar tiempo
)
embedder = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)

def obtener_tipos_demanda_por_abogado(user_id: str) -> List[str]:
    """Obtiene los tipos de demanda disponibles específicos para un abogado."""
    try:
        # Importar aquí para evitar dependencias circulares
        from supabase_integration import supabase_admin
        
        # Buscar el ID del abogado basado en el user_id
        abogado_result = supabase_admin.table("abogados")\
            .select("id")\
            .eq("user_id", user_id)\
            .execute()
        
        if not abogado_result.data or len(abogado_result.data) == 0:
            print(f"⚠️ No se encontró abogado para user_id: {user_id}")
            # Fallback con tipos básicos
            return ["Despido", "Empleados En Blanco", "Empleados En Negro", "Demanda Licencia Medica", "Demanda Solidaridad Laboral"]
        
        abogado_id = abogado_result.data[0]["id"]
        print(f"🔍 Obteniendo tipos para abogado {abogado_id} (user_id: {user_id})")
        
        # Obtener categorías específicas del abogado
        response = supabase_admin.table('categorias_demandas')\
            .select('nombre, id, descripcion')\
            .eq('user_id', user_id)\
            .eq('activo', True)\
            .execute()
        
        if not response.data:
            print(f"⚠️ No hay categorías para el abogado {abogado_id}")
            # Crear categorías por defecto
            return crear_categorias_por_defecto(user_id)
        
        # Verificar cuáles categorías tienen documentos procesados
        categorias_con_documentos = []
        categorias_sin_documentos = []
        
        for categoria in response.data:
            docs_response = supabase_admin.table('documentos_entrenamiento')\
                .select('id')\
                .eq('abogado_id', abogado_id)\
                .eq('categoria_id', categoria['id'])\
                .eq('estado_procesamiento', 'completado')\
                .execute()
            
            if docs_response.data and len(docs_response.data) > 0:
                categorias_con_documentos.append(categoria['nombre'])
            else:
                categorias_sin_documentos.append(categoria['nombre'])
        
        # Priorizar categorías con documentos, pero incluir todas
        tipos_finales = categorias_con_documentos + categorias_sin_documentos
        
        if tipos_finales:
            print(f"✅ Tipos encontrados para abogado {abogado_id}: {tipos_finales}")
            print(f"   📄 Con documentos: {categorias_con_documentos}")
            print(f"   📁 Sin documentos: {categorias_sin_documentos}")
            return tipos_finales
        else:
            print(f"⚠️ No se encontraron tipos para el abogado {abogado_id}")
            return crear_categorias_por_defecto(user_id)
            
    except Exception as e:
        print(f"❌ Error obteniendo tipos por abogado: {e}")
        # Fallback con tipos básicos
        return ["Despido", "Empleados En Blanco", "Empleados En Negro", "Demanda Licencia Medica", "Demanda Solidaridad Laboral"]

def crear_categorias_por_defecto(user_id: str) -> List[str]:
    """Crea categorías por defecto para un abogado nuevo."""
    try:
        from backend.core.category_manager import CategoryManager
        
        category_manager = CategoryManager()
        categorias_creadas = category_manager.create_default_categories_for_user(user_id)
        
        if categorias_creadas:
            nombres = [cat['nombre'] for cat in categorias_creadas]
            print(f"✅ Categorías por defecto creadas: {nombres}")
            return nombres
        else:
            print("⚠️ No se pudieron crear categorías por defecto")
            return ["Despido", "Empleados En Blanco", "Empleados En Negro"]
            
    except Exception as e:
        print(f"❌ Error creando categorías por defecto: {e}")
        return ["Despido", "Empleados En Blanco", "Empleados En Negro"]

# Mantener función original para compatibilidad
def obtener_tipos_demanda_disponibles() -> List[str]:
    """Obtiene los tipos de demanda disponibles (función legacy para compatibilidad)."""
    print("⚠️ Usando función legacy - se recomienda usar obtener_tipos_demanda_por_abogado()")
    return ["Despido", "Empleados En Blanco", "Empleados En Negro", "Demanda Licencia Medica", "Demanda Solidaridad Laboral"]

def mapear_nombre_carpeta(nombre_carpeta: str) -> str:
    """Mapea nombres de carpetas a nombres legibles para el usuario."""
    mapeos = {
        "empleados_blanco": "Empleados En Blanco",
        "empleados_negro": "Empleados En Negro", 
        "empleados_blanco_negro": "Empleados Blanco Negro",
        "demanda_licencia_medica": "Demanda Licencia Medica",
        "demanda_solidaridad_laboral": "Demanda Solidaridad Laboral"
    }
    
    return mapeos.get(nombre_carpeta.lower(), nombre_carpeta.replace("_", " ").title())

def mapear_tipo_a_coleccion(tipo_demanda: str) -> str:
    """Mapea tipos de demanda legibles a nombres de colección en Qdrant."""
    mapeo_inverso = {
        "Empleados En Blanco": "empleados_en_blanco",
        "Empleados En Negro": "empleados_en_negro",
        "Empleados Blanco Negro": "empleados_blanco_negro", 
        "Demanda Licencia Medica": "demanda_licencia_medica",
        "Demanda Solidaridad Laboral": "demanda_solidaridad_laboral"
    }
    
    return mapeo_inverso.get(tipo_demanda, tipo_demanda.lower().replace(" ", "_"))

def buscar_contexto_legal_enriquecido(user_id: str, query: str, tipo_demanda: str = None, info_documentos: Dict = None) -> str:
    """Busca contexto legal enriquecido con anotaciones de experto en colección personal."""
    
    try:
        # Usar DocumentProcessor para búsqueda con anotaciones
        from backend.core.document_processor import DocumentProcessor
        
        doc_processor = DocumentProcessor()
        enhanced_context = doc_processor.get_enhanced_legal_context(
            user_id=user_id,
            query_text=query,
            tipo_demanda=tipo_demanda,
            limit=3
        )
        
        # Construir contexto estructurado para IA
        context_parts = []
        
        if enhanced_context.get("contexto_textual"):
            context_parts.append(f"CONTEXTO BASE: {enhanced_context['contexto_textual']}")
        
        if enhanced_context.get("estrategias_expertas"):
            estrategias_text = " | ".join(enhanced_context["estrategias_expertas"])
            context_parts.append(f"ESTRATEGIAS EXPERTAS: {estrategias_text}")
        
        if enhanced_context.get("precedentes_identificados"):
            precedentes_text = " | ".join(enhanced_context["precedentes_identificados"])
            context_parts.append(f"PRECEDENTES APLICABLES: {precedentes_text}")
        
        if enhanced_context.get("problemas_comunes"):
            problemas_text = " | ".join(enhanced_context["problemas_comunes"])
            context_parts.append(f"PROBLEMAS DETECTADOS: {problemas_text}")
        
        if enhanced_context.get("insights_adicionales"):
            insights_text = " | ".join(enhanced_context["insights_adicionales"])
            context_parts.append(f"INSIGHTS ADICIONALES: {insights_text}")
        
        # NUEVO: Enriquecer con información de documentos procesados
        if info_documentos and info_documentos.get("transcripcion_completa"):
            context_parts.append(f"DOCUMENTOS PROCESADOS: {info_documentos['transcripcion_completa'][:1000]}")
            
            # Agregar información estructurada de documentos
            if info_documentos.get("personas_identificadas"):
                personas_text = ", ".join(info_documentos["personas_identificadas"][:5])
                context_parts.append(f"PERSONAS EN DOCUMENTOS: {personas_text}")
            
            if info_documentos.get("empresas_identificadas"):
                empresas_text = ", ".join(info_documentos["empresas_identificadas"][:3])
                context_parts.append(f"EMPRESAS EN DOCUMENTOS: {empresas_text}")
            
            if info_documentos.get("fechas_importantes"):
                fechas_text = ", ".join(info_documentos["fechas_importantes"][:5])
                context_parts.append(f"FECHAS EN DOCUMENTOS: {fechas_text}")
            
            if info_documentos.get("montos_encontrados"):
                montos_text = ", ".join(info_documentos["montos_encontrados"][:3])
                context_parts.append(f"MONTOS EN DOCUMENTOS: {montos_text}")
        
        # Añadir información de la fuente
        if enhanced_context.get("documentos_fuente"):
            stats_text = f"[Fuentes: {enhanced_context['documentos_fuente']} docs, {enhanced_context.get('documentos_con_anotaciones', 0)} con anotaciones expertas]"
            context_parts.append(stats_text)
        
        # Si tenemos algún contexto, devolverlo
        if context_parts:
            return " || ".join(context_parts)
        else:
            print("⚠️ Contexto enriquecido vacío, usando fallback")
            raise Exception("Contexto enriquecido vacío")
        
    except Exception as e:
        print(f"❌ Error en búsqueda enriquecida, usando fallback: {e}")
        # Usar fallback más simple
        try:
            return buscar_contexto_usuario_basico(user_id, query, tipo_demanda, 3)
        except Exception as e2:
            print(f"❌ Error en fallback básico: {e2}")
            # Último fallback: contexto predefinido
            return obtener_contexto_predefinido(tipo_demanda or "Derecho Laboral")

def buscar_contexto_usuario_basico(user_id: str, query: str, tipo_demanda: str = None, top_k: int = 3) -> str:
    """Busca contexto en la colección personal del usuario (fallback cuando falla búsqueda enriquecida)."""
    
    try:
        # Usar DocumentProcessor para acceder a la colección personal
        from backend.core.document_processor import DocumentProcessor
        
        doc_processor = DocumentProcessor()
        collection_name = doc_processor.get_user_collection_name(user_id)
        
        print(f"🔍 Buscando en colección personal: {collection_name}")
        
        # Verificar si la colección existe
        if not doc_processor.collection_exists(collection_name):
            print(f"⚠️ Colección personal {collection_name} no existe")
            return obtener_contexto_predefinido(tipo_demanda)
        
        # Buscar documentos relevantes en la colección personal
        resultados = doc_processor.search_similar_documents(
            user_id=user_id,
            query_text=query,
            categoria_id=None,  # Buscar en todas las categorías del usuario
            limit=top_k
        )
        
        if not resultados:
            print(f"⚠️ No se encontraron documentos relevantes en colección personal")
            return obtener_contexto_predefinido(tipo_demanda)
        
        # Construir contexto desde los documentos encontrados
        contexto_parts = []
        for i, resultado in enumerate(resultados):
            documento = resultado.get('documento', {})
            payload = documento.get('payload', {})
            
            # Extraer texto relevante
            texto = resultado.get('texto_relevante', '') or payload.get('texto_completo', '')[:300]
            filename = payload.get('filename', f'Documento {i+1}')
            
            if texto:
                parte_contexto = f"DOCUMENTO {i+1} ({filename}): {texto}"
                if len(parte_contexto) > 400:  # Limitar longitud
                    parte_contexto = parte_contexto[:400] + "..."
                contexto_parts.append(parte_contexto)
        
        if contexto_parts:
            contexto_personal = " || ".join(contexto_parts)
            print(f"✅ Contexto extraído de {len(contexto_parts)} documentos personales")
            return f"CONTEXTO DE TUS DOCUMENTOS PERSONALES: {contexto_personal}"
        else:
            return obtener_contexto_predefinido(tipo_demanda)
            
    except Exception as e:
        print(f"❌ Error buscando en colección personal: {e}")
        return obtener_contexto_predefinido(tipo_demanda)

def buscar_contexto_legal_basico(tipo: str, query: str, top_k: int = 2) -> str:
    """Búsqueda básica en colecciones públicas (fallback)."""
    # Mapear el tipo legible al nombre de colección
    nombre_coleccion = mapear_tipo_a_coleccion(tipo)
    collection = f"{COLLECTION_PREFIX}{nombre_coleccion}"
    
    try:
        emb = embedder.embed_query(query)
        search = qdrant.search(
            collection_name=collection, 
            query_vector=emb, 
            limit=top_k, 
            with_payload=True
        )
        
        contexto_parts = []
        for i, resultado in enumerate(search):
            payload = resultado.payload
            # Solo incluir partes clave y resumidas
            hechos = payload.get('hechos', '')[:300] + "..." if len(payload.get('hechos', '')) > 300 else payload.get('hechos', '')
            derecho = payload.get('derecho', '')[:200] + "..." if len(payload.get('derecho', '')) > 200 else payload.get('derecho', '')
            
            parte_contexto = f"CASO {i+1}: HECHOS: {hechos} | DERECHO: {derecho}"
            contexto_parts.append(parte_contexto)
        
        return " | ".join(contexto_parts) if contexto_parts else obtener_contexto_predefinido(tipo)
    except Exception as e:
        print(f"Error buscando contexto para tipo '{tipo}' (colección: {collection}): {e}")
        return obtener_contexto_predefinido(tipo)

def obtener_contexto_predefinido(tipo_demanda: str) -> str:
    """Proporciona contexto legal predefinido cuando no hay casos en la base de datos."""
    
    contextos_predefinidos = {
        "Derecho Laboral": """
        CONTEXTO LEGAL PREDEFINIDO - DERECHO LABORAL:
        
        HECHOS TÍPICOS: Despido injustificado de empleado en relación de dependencia. Trabajador con antigüedad que fue desvinculado sin causa justa invocada por el empleador.
        
        DERECHO APLICABLE: 
        - Ley de Contrato de Trabajo 20.744 - Art. 245: Indemnización por despido sin justa causa
        - LCT Art. 232: Preaviso 
        - LCT Art. 233: Integración del mes de despido
        - LCT Art. 258: Indemnización por antigüedad
        
        JURISPRUDENCIA: "El despido sin causa es una facultad del empleador que debe indemnizar según los arts. 232, 233 y 245 de la LCT" (CNAT).
        
        PETITORIO TÍPICO: Indemnización por despido (Art. 245), preaviso (Art. 232), integración del mes de despido (Art. 233), diferencias salariales si las hubiere.
        
        PRUEBAS HABITUALES: Recibos de sueldo, telegrama de despido, certificado de servicios, testimonios de compañeros de trabajo.
        """,
        
        "Empleados En Blanco": """
        CONTEXTO LEGAL PREDEFINIDO - EMPLEADOS EN BLANCO:
        
        HECHOS TÍPICOS: Trabajador registrado que sufrió despido sin causa. Relación laboral documentada con aportes y contribuciones al día.
        
        DERECHO APLICABLE:
        - LCT Art. 245: Indemnización por despido sin justa causa equivalente a un mes de sueldo por año de antigüedad
        - LCT Art. 232/233: Preaviso e integración
        - Ley 25.323: Incremento indemnizatorio por falta de pago en término
        
        MONTO INDEMNIZATORIO: Un mes de sueldo por año de antigüedad o fracción mayor a 3 meses.
        """,
        
        "Empleados En Negro": """
        CONTEXTO LEGAL PREDEFINIDO - TRABAJADORES NO REGISTRADOS:
        
        HECHOS TÍPICOS: Trabajador no registrado o parcialmente registrado. Prestación de servicios sin alta tempestiva en AFIP.
        
        DERECHO APLICABLE:
        - Ley 24.013 Art. 8, 9 y 10: Indemnizaciones por falta de registración
        - Ley 25.323: Incrementos indemnizatorios
        - LCT Art. 245: Indemnización común por despido
        
        INDEMNIZACIONES ESPECIALES: Doble indemnización del Art. 245, plus del 25% del Art. 9, indemnización sustitutiva del preaviso incrementada.
        """
    }
    
    # Buscar contexto específico o usar genérico
    contexto = contextos_predefinidos.get(tipo_demanda)
    if contexto:
        return contexto.strip()
    
    # Contexto genérico para cualquier tipo
    return """
    CONTEXTO LEGAL GENÉRICO - DERECHO LABORAL:
    
    PRINCIPIOS FUNDAMENTALES:
    - Principio protectorio del trabajador (Art. 9 LCT)
    - Indubio pro operario en caso de duda
    - Primacía de la realidad sobre las formas
    
    BASE LEGAL COMÚN:
    - Ley de Contrato de Trabajo 20.744
    - Ley Nacional de Empleo 24.013  
    - Ley de Riesgos del Trabajo 24.557
    - Código Civil y Comercial (subsidiario)
    
    ESTRUCTURA DE DEMANDA:
    1. Hechos cronológicos y precisos
    2. Derecho aplicable con citas normativas
    3. Petitorio claro y fundado
    4. Ofrecimiento de prueba específico
    """

# Mantener función original para compatibilidad
def buscar_contexto_legal(tipo: str, query: str, top_k: int = 2) -> str:
    """Busca contexto legal relevante en la base de datos vectorial (función legacy)."""
    return buscar_contexto_legal_basico(tipo, query, top_k)

def obtener_informacion_documentos_sincrona(session_id: str) -> Dict:
    """Obtiene información consolidada de documentos de forma síncrona."""
    try:
        print(f"📄 Obteniendo información consolidada de documentos para session_id: {session_id}")
        from supabase_integration import supabase_admin
        
        # Obtener consolidado usando función SQL directamente
        response = supabase_admin.rpc('get_documentos_chat_consolidado', {
            'p_session_id': session_id
        }).execute()
        
        if response.data and len(response.data) > 0:
            consolidado = response.data[0]
            informacion_documentos = {
                'transcripcion_completa': consolidado.get('transcripcion_completa', ''),
                'personas_identificadas': consolidado.get('personas_identificadas', []),
                'empresas_identificadas': consolidado.get('empresas_identificadas', []),
                'fechas_importantes': consolidado.get('fechas_importantes', []),
                'montos_encontrados': consolidado.get('montos_encontrados', []),
                'datos_contacto': consolidado.get('datos_contacto', {})
            }
            print(f"✅ Información de documentos obtenida: {len(informacion_documentos['personas_identificadas'])} personas, {len(informacion_documentos['empresas_identificadas'])} empresas")
            return informacion_documentos
        else:
            print(f"⚠️ No se encontró información consolidada para session_id: {session_id}")
            return {}
            
    except Exception as e:
        print(f"⚠️ Error obteniendo información de documentos: {e}")
        # Continuar sin información de documentos
        return {}

async def generar_demanda(tipo_demanda: str, datos_cliente: Dict, hechos_adicionales: str = "", notas_abogado: str = "", user_id: str = None, session_id: str = None) -> Dict:
    """Genera una demanda legal usando IA con contexto enriquecido."""
    
    try:
        print(f"🚀 Iniciando generación de demanda: {tipo_demanda}")
        print(f"👤 Cliente: {datos_cliente.get('nombre_completo', 'No especificado')}")
        
        # NUEVO: Obtener información de documentos procesados
        info_documentos = {}
        if session_id:
            try:
                info_documentos = obtener_informacion_documentos_sincrona(session_id)
                print(f"📄 Documentos procesados: {len(info_documentos.get('transcripcion_completa', ''))} caracteres")
            except Exception as e:
                print(f"⚠️ Error obteniendo documentos: {e}")
        
        # Construir query para búsqueda de contexto
        query = f"{tipo_demanda} {hechos_adicionales}"
        
        # NUEVO: Buscar contexto legal enriquecido con información de documentos
        contexto_legal = buscar_contexto_legal_enriquecido(
            user_id=user_id or "default", 
            query=query, 
            tipo_demanda=tipo_demanda,
            info_documentos=info_documentos
        )
        
        # Construir prompt completo con toda la información
        prompt = f"""
Eres un abogado experto en derecho laboral argentino. Genera una demanda profesional basándote en la siguiente información:

TIPO DE DEMANDA: {tipo_demanda}

DATOS DEL CLIENTE:
- Nombre: {datos_cliente.get('nombre_completo', 'No especificado')}
- DNI: {datos_cliente.get('dni', 'No especificado')}
- Domicilio: {datos_cliente.get('domicilio', 'No especificado')}
- Teléfono: {datos_cliente.get('telefono', 'No especificado')}
- Email: {datos_cliente.get('email', 'No especificado')}
- Ocupación: {datos_cliente.get('ocupacion', 'No especificado')}

HECHOS ADICIONALES:
{hechos_adicionales}

NOTAS DEL ABOGADO:
{notas_abogado}

CONTEXTO LEGAL Y JURISPRUDENCIA:
{contexto_legal}

INFORMACIÓN DE DOCUMENTOS PROCESADOS:
{info_documentos.get('transcripcion_completa', 'No hay documentos procesados')}

INSTRUCCIONES:
1. Genera una demanda completa y profesional
2. Incluye todos los elementos legales necesarios
3. Usa la información de documentos procesados cuando sea relevante
4. Cita artículos específicos de la LCT y jurisprudencia
5. Estructura la demanda correctamente (hechos, derecho, petitorio, prueba)
6. Adapta el contenido al tipo de demanda específico

Genera la demanda completa:
"""

        # Generar la demanda con IA
        response = llm.invoke(prompt)
        texto_demanda = response.content
        
        # Crear documento Word
        archivo_docx = crear_documento_word(texto_demanda, datos_cliente, tipo_demanda)
        
        # Metadatos enriquecidos
        metadatos = {
            "tipo_demanda": tipo_demanda,
            "datos_cliente": datos_cliente,
            "hechos_adicionales": hechos_adicionales,
            "notas_abogado": notas_abogado,
            "documentos_procesados": len(info_documentos.get('transcripcion_completa', '')),
            "personas_identificadas": info_documentos.get('personas_identificadas', []),
            "empresas_identificadas": info_documentos.get('empresas_identificadas', []),
            "fechas_importantes": info_documentos.get('fechas_importantes', []),
            "montos_encontrados": info_documentos.get('montos_encontrados', []),
            "contexto_legal_usado": len(contexto_legal),
            "fecha_generacion": datetime.now().isoformat()
        }
        
        print(f"✅ Demanda generada exitosamente")
        print(f"📄 Archivo creado: {archivo_docx}")
        print(f"📊 Metadatos: {len(metadatos)} campos")
        
        return {
            "texto_demanda": texto_demanda,
            "archivo_docx": archivo_docx,
            "metadatos": metadatos,
            "exito": True
        }
        
    except Exception as e:
        print(f"❌ Error generando demanda: {e}")
        return {
            "texto_demanda": f"Error generando demanda: {str(e)}",
            "archivo_docx": None,
            "metadatos": {"error": str(e)},
            "exito": False
        }

def generar_demanda_sincrona(tipo_demanda: str, datos_cliente: Dict, hechos_adicionales: str = "", notas_abogado: str = "", user_id: str = None, session_id: str = None) -> Dict:
    """Versión síncrona de generar_demanda para compatibilidad con código existente."""
    import asyncio
    
    try:
        # Verificar si ya hay un event loop corriendo
        try:
            loop = asyncio.get_running_loop()
            # Si hay un loop corriendo, crear uno nuevo
            import concurrent.futures
            with concurrent.futures.ThreadPoolExecutor() as executor:
                future = executor.submit(asyncio.run, generar_demanda(
                    tipo_demanda=tipo_demanda,
                    datos_cliente=datos_cliente,
                    hechos_adicionales=hechos_adicionales,
                    notas_abogado=notas_abogado,
                    user_id=user_id,
                    session_id=session_id
                ))
                return future.result()
        except RuntimeError:
            # No hay loop corriendo, usar asyncio.run
            return asyncio.run(generar_demanda(
                tipo_demanda=tipo_demanda,
                datos_cliente=datos_cliente,
                hechos_adicionales=hechos_adicionales,
                notas_abogado=notas_abogado,
                user_id=user_id,
                session_id=session_id
            ))
    except Exception as e:
        raise Exception(f"Error en generación síncrona: {str(e)}")

def crear_documento_word(texto_demanda: str, datos_cliente: Dict, tipo_demanda: str) -> str:
    """Crea un documento Word con la demanda generada - tipografía completamente en negro."""
    try:
        from docx.shared import RGBColor
        from docx.enum.text import WD_COLOR_INDEX
        
        # Crear documento
        doc = Document()
        
        # Configurar título en negro
        titulo = doc.add_heading(f'DEMANDA - {tipo_demanda.upper()}', 0)
        titulo.alignment = 1  # Centrado
        # Asegurar color negro en título
        for run in titulo.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # Negro
        
        # Agregar información del cliente
        encabezado_cliente = doc.add_heading('DATOS DEL CLIENTE', level=1)
        # Asegurar color negro en encabezado
        for run in encabezado_cliente.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        tabla_cliente = doc.add_table(rows=0, cols=2)
        tabla_cliente.style = 'Table Grid'
        
        for campo, valor in datos_cliente.items():
            if valor:
                row = tabla_cliente.add_row()
                # Configurar texto en negro para la tabla
                cell_campo = row.cells[0]
                cell_valor = row.cells[1]
                
                cell_campo.text = campo.replace('_', ' ').title()
                cell_valor.text = str(valor)
                
                # Asegurar color negro en celdas
                for paragraph in cell_campo.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                for paragraph in cell_valor.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Agregar salto de página
        doc.add_page_break()
        
        # Agregar texto de la demanda
        encabezado_demanda = doc.add_heading('DEMANDA', level=1)
        # Asegurar color negro en encabezado
        for run in encabezado_demanda.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Dividir el texto en párrafos y procesar
        paragrafos = texto_demanda.split('\n\n')
        for parrafo in paragrafos:
            if parrafo.strip():
                # Verificar si es un encabezado de sección
                es_encabezado = any(keyword in parrafo.upper() for keyword in ['HECHOS', 'DERECHO', 'PETITORIO', 'PRUEBA'])
                
                if es_encabezado:
                    # Crear como encabezado pero en negro
                    p = doc.add_heading(parrafo.strip(), level=2)
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                else:
                    # Crear como párrafo normal en negro
                    p = doc.add_paragraph(parrafo.strip())
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Guardar en archivo temporal
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        
        return temp_file.name
        
    except Exception as e:
        raise Exception(f"Error creando documento Word: {str(e)}")

def validar_datos_cliente(datos: Dict) -> Dict:
    """Valida y formatea los datos del cliente."""
    campos_requeridos = ['nombre_completo', 'dni', 'domicilio']
    campos_faltantes = [campo for campo in campos_requeridos if not datos.get(campo)]
    
    if campos_faltantes:
        raise ValueError(f"Faltan campos obligatorios: {', '.join(campos_faltantes)}")
    
    return datos

# Ejemplo de uso:
if __name__ == "__main__":
    datos = {
        "nombre": "Romina Gabriela Ruth Vaca",
        "dni": "32501402",
        "domicilio": "Itaqui 2099, CABA",
        "fecha_nacimiento": "27/10/1986",
        "email": "v.romina7529@gmail.com",
        "estado_civil": "Soltera",
        "motivo": "despido mientras estaba con licencia médica"
    }
    resultado = generar_demanda("Demanda con licencia medica", datos)
    print(resultado)